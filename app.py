"""
XGL South Africa — Meeting Minutes DOCX Microservice
Flask service: receives JSON → generates DOCX from company template → returns binary file
Deploy on Render.com (free tier) or any Python host
"""

import os
import io
import copy
import logging
import uuid
from datetime import datetime
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

app = Flask(__name__)
logging.basicConfig(level=logging.INFO)

# Template file path — should sit next to app.py
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_heading(doc, text: str, level: int = 1, color: str = None):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    elif level == 2:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # medium blue
    elif level == 3:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    run.font.name = 'Assistant'
    return p

def add_bilingual_para(doc, en_text: str, zh_text: str,
                        prefix: str = '', bold_prefix: bool = True,
                        indent: bool = False):
    """Add EN line followed by ZH line."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    if prefix:
        r = p.add_run(prefix + ' ')
        r.font.bold = bold_prefix
        r.font.name = 'Assistant'
        r.font.size = Pt(10)
    r_en = p.add_run(en_text)
    r_en.font.name = 'Assistant'
    r_en.font.size = Pt(10)

    p2 = doc.add_paragraph()
    if indent:
        p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(6)
    r_zh = p2.add_run(zh_text)
    r_zh.font.name = 'Assistant'
    r_zh.font.size = Pt(10)
    r_zh.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p, p2

def add_divider(doc):
    """Add a thin horizontal rule using paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_meta_table(doc, meeting_info: dict):
    """Add a 2-column info table for meeting metadata."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    set_cell_borders(tbl)

    def add_row(label, value):
        row = tbl.add_row()
        # label cell
        lc = row.cells[0]
        lc.width = Inches(1.4)
        set_cell_bg(lc, 'EBF3FB')
        rp = lc.paragraphs[0]
        r = rp.add_run(label)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Assistant'
        # value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value) if value else '—')
        vr.font.size = Pt(9)
        vr.font.name = 'Assistant'

    add_row('Title / 标题 (EN)', meeting_info.get('title_en', ''))
    add_row('Title / 标题 (ZH)', meeting_info.get('title_zh', ''))
    add_row('Date / 日期', meeting_info.get('date', ''))
    add_row('Location / 地点', meeting_info.get('location', ''))

    participants = meeting_info.get('participants', [])
    if isinstance(participants, list):
        participants_str = '\n'.join(participants)
    else:
        participants_str = str(participants)
    add_row('Participants / 与会人员', participants_str)

    doc.add_paragraph()  # spacer

def add_decisions_section(doc, decisions: list):
    """Add key decisions section."""
    if not decisions:
        return
    add_heading(doc, '2. Key Decisions / 主要决议', level=1)
    for i, item in enumerate(decisions, 1):
        en = item.get('en', '')
        zh = item.get('zh', '')
        speaker = item.get('speaker', '')
        time_range = item.get('evidence_time_range', '')

        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(en)
        r.font.size = Pt(10)
        r.font.name = 'Assistant'

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.3)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(zh)
        r2.font.size = Pt(10)
        r2.font.name = 'Assistant'
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if speaker or time_range:
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Inches(0.3)
            p3.paragraph_format.space_after = Pt(8)
            meta_parts = []
            if speaker:
                meta_parts.append(f'Speaker: {speaker}')
            if time_range:
                meta_parts.append(f'Ref: {time_range}')
            r3 = p3.add_run('  '.join(meta_parts))
            r3.font.size = Pt(8)
            r3.font.name = 'Assistant'
            r3.font.italic = True
            r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_action_items_table(doc, action_items: list):
    """Add action items as a formatted table."""
    if not action_items:
        return
    add_heading(doc, '3. Action Items / 行动事项', level=1)

    headers = ['Task (EN/ZH)', 'Owner', 'Deadline', 'Speaker', 'Evidence']
    col_widths = [Inches(2.8), Inches(1.0), Inches(0.9), Inches(0.9), Inches(0.8)]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    set_cell_borders(tbl)

    # header row
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        set_cell_bg(cell, '1F497D')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8)
        r.font.name = 'Assistant'
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in action_items:
        row = tbl.add_row()
        task_en = item.get('task_en', '')
        task`zh = item.get('task_zh', '')

        # Task cell (EN + ZH)
        tc = row.cells[0]
        p_en = tc.paragraphs[0]
        r_en = p_en.add_run(task_en)
        r_en.font.size = Pt(8)
        r_en.font.name = 'Assistant'
        p_zh = tc.add_paragraph()
        r_zh = p_zh.add_run(task_zh)
        r_zh.font.size = Pt(8)
        r_zh.font.name = 'Assistant'
        r_zh.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        def fill_cell(c, text, center=False):
            p = c.paragraphs[0]
            r = p.add_run(str(text) if text else '—')
            r.font.size = Pt(8)
            r.font.name = 'Assistant'
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fill_cell(row.cells[1], item.get('owner', ''))
        fill_cell(row.cells[2], item.get('deadline', ''), center=True)
        fill_cell(row.cells[3], item.get('speaker', ''))
        fill_cell(row.cells[4], item.get('evidence_time_range', ''))

    doc.add_paragraph()  # spacer

def add_pending_section(doc, pending: list):
    """Add pending matters section."""
    if not pending:
        return
    add_heading(doc, '4. Pending Matters / 待处理事项', level=1)
    for item in pending:
        en = item.get('en', '')
        zh = item.get('zh', '')
        speaker = item.get('speaker', '')
        time_range = item.get('evidence_time_range', '')
        add_bilingual_para(doc, en, zh, prefix='•', indent=True)
        if speaker or time_range:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.6)
            p.paragraph_format.space_after = Pt(6)
            parts = []
            if speaker:
                parts.append(f'Speaker: {speaker}')
            if time_range:
                parts.append(f'Ref: {time_range}')
            r = p.add_run('  '.join(parts))
            r.font.size = Pt(8)
            r.font.name = 'Assistant'
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_transcript_appendix(doc, full_transcript: list):
    """Add full bilingual transcript as appendix table."""
    if not full_transcript:
        return

    doc.add_page_break()
    add_heading(doc, 'Appendix: Full Bilingual Transcript / 附录：完整双语转录', level=1)

    headers = ['Speaker', 'Time', 'English', '中文']
    col_widths = [Inches(0.9), Inches(0.7), Inches(2.5), Inches(2.5)]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    set_cell_borders(tbl)

    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        set_cell_bg(cell, '2E74B5')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8)
        r.font.name = 'Assistant'
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for entry in full_transcript:
        row = tbl.add_row()
        def fill(c, text, wrap=True):
            p = c.paragraphs[0]
            r = p.add_run(str(text) if text else '')
            r.font.size = Pt(7.5)
            r.font.name = 'Assistant'

        fill(row.cells[0], entry.get('speaker', ''))
        fill(row.cells[1], entry.get('start_time', ''))
        fill(row.cells[2], entry.get('en_text', ''))
        fill(row.cells[3], entry.get('zh_text', ''))


# ── main document builder ───────────────────────────────────────────────────

def build_minutes_docx(meeting_info: dict, minutes: dict, full_transcript: list) -> bytes:
    """Build the DOCX from the template and return as bytes."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")

    try:
        doc = Document(TEMPLATE_PATH)
    except Exception:
        # Template corrupted or missing — fall back to blank document
        doc = Document()
        doc.add_paragraph()  # ensure at least one paragraph exists

    # Clear existing empty paragraphs in body (keep 1 for structure)
    for p in doc.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    # ── Cover / Title block ─────────────────────────────────────────────────
    title_p = doc.paragraphs[0]
    title_p.clear()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run('MEETING MINUTES / 会议记录')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.name = 'Assistant'
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle_p.add_run('CHEETAH CHROME SOUTH AFRICA (PTY) LTD — XGL')
    r2.font.size = Pt(10)
    r2.font.name = 'Assistant'
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    generated_p = doc.add_paragraph()
    generated_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = generated_p.add_run(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')
    r3.font.size = Pt(8)
    r3.font.name = 'Assistant'
    r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    add_divider(doc)

    # ── Section 1: Meeting Info ─────────────────────────────────────────────
    add_heading(doc, '1. Meeting Information / 会议信息', level=1)
    add_meta_table(doc, meeting_info)

    # ── Section 2: Overview ─────────────────────────────────────────────────
    overview = minutes.get('overview', {})
    if overview:
        add_heading(doc, '2. Overview / 概述', level=1)
        add_bilingual_para(doc, overview.get('en', ''), overview.get('zh', ''))
        doc.add_paragraph()

    # ── Section 3: Key Decisions ─────────────────────────────────────────────
    decisions = minutes.get('key_decisions', [])
    # re-number sections since overview might shift numbering
    add_heading(doc, '3. Key Decisions / 主要决议', level=1)
    if decisions:
        for i, item in enumerate(decisions, 1):
            en = item.get('en', '')
            zh = item.get('zh', '')
            speaker = item.get('speaker', '')
            time_range = item.get('evidence_time_range', '')

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_num = p.add_run(f'{i}. ')
            r_num.font.bold = True
            r_num.font.size = Pt(10)
            r_num.font.name = 'Assistant'
            r_en = p.add_run(en)
            r_en.font.size = Pt(10)
            r_en.font.name = 'Assistant'

            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.paragraph_format.space_after = Pt(2)
            r_zh = p2.add_run(zh)
            r_zh.font.size = Pt(10)
            r_zh.font.name = 'Assistant'
            r_zh.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            if speaker or time_range:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.3)
                p3.paragraph_format.space_after = Pt(8)
                parts = []
                if speaker:
                    parts.append(f'Speaker: {speaker}')
                if time_range:
                    parts.append(f'Ref: {time_range}')
                r_meta = p3.add_run('  '.join(parts))
                r_meta.font.size = Pt(8)
                r_meta.font.italic = True
                r_meta.font.name = 'Assistant'
                r_meta.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        p = doc.add_paragraph()
        p.add_run('No key decisions recorded. / 无主要决议。').font.size = Pt(10)

    # ── Section 4: Action Items ─────────────────────────────────────────────
    action_items = minutes.get('action_items', [])
    add_heading(doc, '4. Action Items / 行动事项', level=1)
    if action_items:
        add_action_items_table(doc, action_items)
    else:
        p = doc.add_paragraph()
        p.add_run('No action items recorded. / 无行动事项。').font.size = Pt(10)

    # ── Section 5: Pending Matters ──────────────────────────────────────────
    pending = minutes.get('pending_matters', [])
    add_heading(doc, '5. Pending Matters / 待处理事项', level=1)
    if pending:
        add_pending_section(doc, pending)
    else:
        p = doc.add_paragraph()
        p.add_run('No pending matters. / 无待处理事项。').font.size = Pt(10)

    # ── Appendix: Full Transcript ───────────────────────────────────────────
    if full_transcript:
        add_transcript_appendix(doc, full_transcript)

    # ── Return as bytes ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Flask routes ─────────────────────────────────────────────────────────────

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint."""
    template_ok = os.path.exists(TEMPLATE_PATH)
    return jsonify({
        'status': 'ok',
        'template_found': template_ok,
        'template_path': TEMPLATE_PATH
    })

@app.route('/generate-minutes', methods=['POST'])
def generate_minutes():
    """
    POST /generate-minutes
    Body JSON:
    {
      "meeting_info": {
        "title_en": "...", "title_zh": "...",
        "date": "...", "location": "...",
        "participants": ["name1", ...]
      },
      "minutes": {
        "overview": {"en": "...", "zh": "..."},
        "key_decisions": [{"en": "...", "zh": "...", "speaker": "...", "evidence_time_range": "..."}],
        "action_items": [{"task_en": "...", "task_zh": "...", "owner": "...", "deadline": "...", "speaker": "...", "evidence_time_range": "..."}],
        "pending_matters": [{"en": "...", "zh": "...", "speaker": "...", "evidence_time_range": "..."}]
      },
      "full_transcript": [
        {"speaker": "A", "start_time": "00:01:23", "end_time": "00:01:45", "en_text": "...", "zh_text": "..."}
      ]
    }
    Returns: application/vnd.openxmlformats-officedocument.wordprocessingml.document
    """
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400

        meeting_info = data.get('meeting_info', {})
        minutes = data.get('minutes', {})
        full_transcript = data.get('full_transcript', [])

        app.logger.info(f"Generating DOCX for: {meeting_info.get('title_en', 'Untitled')}")

        docx_bytes = build_minutes_docx(meeting_info, minutes, full_transcript)

        # Build a safe filename
        date_str = meeting_info.get('date', datetime.utcnow().strftime('%Y-%m-%d'))
        safe_date = date_str.replace('/', '-').replace(' ', '_')[:10]
        title_slug = meeting_info.get('title_en', 'Meeting')[:30].replace(' ', '_')
        filename = f"MeetingMinutes_{safe_date}_{title_slug}.docx"

        return send_file(
            io.BytesIO(docx_bytes),
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except FileNotFoundError as e:
        app.logger.error(str(e))
        return jsonify({'error': str(e)}), 500
    except Exception as e:
        app.logger.error(f"Error generating DOCX: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


def _cors(resp):
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return resp


@app.route('/upload-file', methods=['POST', 'OPTIONS'])
def upload_file():
    """
    POST /upload-file  (multipart form-data, field name = 'file')
    Stores the file in /tmp and returns a public download URL.
    Supports CORS from any origin so browser pages can call this directly.
    """
    if request.method == 'OPTIONS':
        return _cors(app.make_response(''))

    if 'file' not in request.files:
        return _cors(jsonify({'error': 'No file field in request'})), 400

    f = request.files['file']
    token = str(uuid.uuid4())
    ext = os.path.splitext(f.filename)[1] if f.filename else '.bin'
    stored_name = f"{token}{ext}"
    path = f"/tmp/{stored_name}"
    f.save(path)

    base_url = request.host_url.rstrip('/')
    url = f"{base_url}/serve/{stored_name}"
    app.logger.info(f"File stored: {stored_name} → {url}")
    return _cors(jsonify({'url': url}))


@app.route('/serve/<stored_name>', methods=['GET'])
def serve_file(stored_name):
    """Serve a previously uploaded temp file from /tmp."""
    # Basic security: only allow uuid-style names with common audio/video extensions
    allowed_exts = {'.m4a', '.mp4', '.mp3', '.wav', '.aac', '.ogg', '.flac', '.webm', '.bin'}
    ext = os.path.splitext(stored_name)[1].lower()
    if ext not in allowed_exts:
        return jsonify({'error': 'File type not allowed'}), 403
    path = f"/tmp/{stored_name}"
    if not os.path.exists(path):
        return jsonify({'error': 'File not found or expired'}), 404
    return send_file(path, as_attachment=False)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=False)
